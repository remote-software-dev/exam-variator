import json
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def convert_to_latex_math(text):
    """
    Convert simple math notation to LaTeX format.
    Examples:
    - 3^(2/3) -> 3^{\\frac{2}{3}}
    - 2^3 -> 2^{3}
    - x × y -> x \\times y
    """
    # Convert fractions in exponents: ^(a/b) -> ^{\\frac{a}{b}}
    text = re.sub(r'\^\((\d+)/(\d+)\)', r'^{\\frac{\1}{\2}}', text)
    text = re.sub(r'\^(\d+)/(\d+)', r'^{\\frac{\1}{\2}}', text)
    
    # Convert simple exponents: ^3 -> ^{3}
    text = re.sub(r'\^(\d+)', r'^{\1}', text)
    
    # Convert multiplication symbol
    text = text.replace('×', '\\times ')
    text = text.replace('x', '\\times ')
    
    # Convert division
    text = text.replace('/', ' / ')
    
    return text

def add_math_equation(paragraph, text):
    """Add text with LaTeX-style math formatting."""
    run = paragraph.add_run(text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(12)

def add_question_block(doc, title_text, q_data):
    """Add a question block with proper formatting."""
    # Fix spacing in title
    title_text = re.sub(r'([a-zA-Z])([A-Z])', r'\1 \2', title_text)
    title_text = re.sub(r'(\d)([A-Z])', r'\1 \2', title_text)
    
    doc.add_heading(title_text, level=1)
    
    # Add Question Text with proper spacing
    p = doc.add_paragraph()
    question_text = q_data["question_text"]
    # Fix spacing: add space between words and numbers/symbols
    question_text = re.sub(r'([a-zA-Z])(\d)', r'\1 \2', question_text)
    question_text = re.sub(r'(\d)([a-zA-Z])', r'\1 \2', question_text)
    question_text = re.sub(r'([)])([a-zA-Z])', r'\1 \2', question_text)
    question_text = re.sub(r'([a-zA-Z])([(])', r'\1 \2', question_text)
    
    add_math_equation(p, question_text)
    
    # Add Options (A, B, C, D, E)
    for i, option in enumerate(q_data["options"]):
        p_option = doc.add_paragraph()
        option_text = f"{chr(65+i)}. {option}"
        # Fix spacing in options too
        option_text = re.sub(r'([)])([a-zA-Z])', r'\1 \2', option_text)
        run = p_option.add_run(option_text)
        run.font.name = "Cambria Math"
        run.font.size = Pt(11)

def main():
    # Sample data
    sample_data = {
        "original": {
            "id": "25MATBLGBRLM01SU-000000-0246",
            "question_text": "Bentuk sederhana dari (3^(2/3) × 8^(3/2)) / (2^(5/2) × 9^(5/6)) adalah ....",
            "options": ["1/42", "2/3", "4/3", "6", "12"]
        },
        "variations": {
            "easier": {
                "question_text": "Bentuk sederhana dari 2³ × 4² / 8¹ adalah ....",
                "options": ["2", "4", "8", "16", "32"]
            },
            "harder": {
                "question_text": "Bentuk sederhana dari (27^(4/3) × 16^(3/4)) / (8^(5/3) × 9^(3/2)) adalah ....",
                "options": ["1/9", "2/9", "4/9", "8/9", "1"]
            }
        }
    }

    # Generate the Word Document
    doc = Document()
    
    # Title with proper spacing
    title = doc.add_heading('Bank Soal & Variasi Matematika', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('SMA/MA/SMK - Tahun 2025')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].italic = True
    
    doc.add_paragraph()  # Spacer

    # Add Original Question
    add_question_block(doc, f"Soal Asli (ID: {sample_data['original']['id']})", sample_data['original'])
    
    doc.add_page_break()
    
    # Add Variations
    if "variations" in sample_data:
        if "easier" in sample_data["variations"]:
            add_question_block(doc, "Variasi Lebih Mudah", sample_data["variations"]["easier"])
        
        doc.add_paragraph()  # Spacer
        
        if "harder" in sample_data["variations"]:
            add_question_block(doc, "Variasi Lebih Sulit", sample_data["variations"]["harder"])

    # Save
    output_path = "data/outputs/soal_variasi_final.docx"
    doc.save(output_path)
    print(f"✅ Generated: {output_path}")
    print("\n📝 For TRUE professional equations (rendered fractions, radicals, etc.),")
    print("   we will integrate a LaTeX equation renderer in the next iteration.")

if __name__ == "__main__":
    main()

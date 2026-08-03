import json
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    # 1. Create a sample JSON file to simulate the output from our AI pipeline
    sample_data = {
        "original": {
            "id": "25MATBLGBRLM01SU-000000-0246",
            "question_text": "Bentuk sederhana dari 3^(2/3) x 8^(3/2) / 2^(5/2) x 9^(5/6) adalah ....",
            "options": ["1/42", "2/3", "4/3", "6", "12"]
        },
        "variations": {
            "easier": {
                "question_text": "Bentuk sederhana dari 2^3 x 4^2 / 8^1 adalah ....",
                "options": ["2", "4", "8", "16", "32"]
            },
            "harder": {
                "question_text": "Bentuk sederhana dari (27^(4/3) x 16^(3/4)) / (8^(5/3) x 9^(3/2)) adalah ....",
                "options": ["1/9", "2/9", "4/9", "8/9", "1"]
            }
        }
    }

    json_path = "data/outputs/sample_questions.json"
    os.makedirs(os.path.dirname(json_path), exist_ok=True)
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(sample_data, f, ensure_ascii=False, indent=2)

    # 2. Read the JSON back (simulating the actual app workflow)
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # 3. Generate the Word Document
    doc = Document()
    
    # Title
    title = doc.add_heading('Bank Soal & Variasi Matematika', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Helper function to format a question block
    def add_question_block(doc, title_text, q_data):
        doc.add_heading(title_text, level=1)
        # Add Question Text
        p = doc.add_paragraph()
        p.add_run(q_data["question_text"]).font.size = Pt(12)
        
        # Add Options (A, B, C, D, E)
        for i, option in enumerate(q_data["options"]):
            doc.add_paragraph(f"{chr(65+i)}. {option}", style='List Bullet')

    # Add Original Question
    add_question_block(doc, f"Soal Asli (ID: {data['original']['id']})", data['original'])
    
    # Add Variations
    if "variations" in data:
        if "easier" in data["variations"]:
            add_question_block(doc, "Variasi Lebih Mudah", data["variations"]["easier"])
        if "harder" in data["variations"]:
            add_question_block(doc, "Variasi Lebih Sulit", data["variations"]["harder"])

    # 4. Save the document
    output_path = "data/outputs/soal_variasi.docx"
    doc.save(output_path)
    print(f"✅ Successfully generated Word document: {output_path}")

if __name__ == "__main__":
    main()

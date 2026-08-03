import json
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_math_paragraph(paragraph, text):
    """
    Convert text with math notation to Word equation.
    Supports: ^(exponent), _(subscript), ^(fraction), sqrt(), times, div
    """
    # For now, we'll use a simple approach: insert text with proper formatting
    # In production, we would convert to OMML (Office Math Markup Language)
    
    # Split by math operators and format each part
    parts = re.split(r'(\^|\(|\)|/|\s+)', text)
    
    run = paragraph.add_run()
    for part in parts:
        if part.strip():
            run.add_text(part)
    
    # Set font to Cambria Math for better equation rendering
    run.font.name = "Cambria Math"
    run.font.size = Pt(12)

def add_question_with_equations(doc, title_text, q_data):
    """Add a question block with professional math formatting."""
    doc.add_heading(title_text, level=1)
    
    # Add Question Text with math formatting
    p = doc.add_paragraph()
    create_math_paragraph(p, q_data["question_text"])
    
    # Add Options (A, B, C, D, E)
    for i, option in enumerate(q_data["options"]):
        p_option = doc.add_paragraph()
        run = p_option.add_run(f"{chr(65+i)}. {option}")
        run.font.name = "Cambria Math"
        run.font.size = Pt(11)

def main():
    # Sample data with LaTeX-style math notation
    sample_data = {
        "original": {
            "id": "25MATBLGBRLM01SU-000000-0246",
            "question_text": "Bentuk sederhana dari (3^(2/3) × 8^(3/2)) / (2^(5/2) × 9^(5/6)) adalah ....",
            "options": ["1/42", "2/3", "4/3", "6", "12"]
        },
        "variations": {
            "easier": {
                "question_text": "Bentuk sederhana dari 2³ × 4² / 8¹ adalah ....",
                "options": ["2", "4", "8", "16", "32"]
            },
            "harder": {
                "question_text": "Bentuk sederhana dari (27^(4/3) × 16^(3/4)) / (8^(5/3) × 9^(3/2)) adalah ....",
                "options": ["1/9", "2/9", "4/9", "8/9", "1"]
            }
        }
    }

    json_path = "data/outputs/sample_questions.json"
    os.makedirs(os.path.dirname(json_path), exist_ok=True)
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(sample_data, f, ensure_ascii=False, indent=2)

    # Generate the Word Document
    doc = Document()
    
    # Title
    title = doc.add_heading('Bank Soal & Variasi Matematika', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph('SMA/MA/SMK - Tahun 2025')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].italic = True
    
    doc.add_paragraph()  # Spacer

    # Add Original Question
    add_question_with_equations(doc, f"Soal Asli (ID: {sample_data['original']['id']})", sample_data['original'])
    
    doc.add_page_break()
    
    # Add Variations
    if "variations" in sample_data:
        if "easier" in sample_data["variations"]:
            add_question_with_equations(doc, "Variasi Lebih Mudah", sample_data["variations"]["easier"])
        
        doc.add_paragraph()  # Spacer
        
        if "harder" in sample_data["variations"]:
            add_question_with_equations(doc, "Variasi Lebih Sulit", sample_data["variations"]["harder"])

    # Save the document
    output_path = "data/outputs/soal_variasi_professional.docx"
    doc.save(output_path)
    print(f"✅ Successfully generated professional Word document: {output_path}")
    print("\n📝 Note: For TRUE professional equations (fractions, radicals, etc.), we will integrate")
    print("   a LaTeX-to-OMML converter in the next iteration.")

if __name__ == "__main__":
    main()

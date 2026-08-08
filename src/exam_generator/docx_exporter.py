"""Markdown-aware DOCX exporter for the exam generator.

Renders a professional Word document from the structured question data with
Markdown formatting (bold, italics, lists) and LaTeX math rendered as native
Word equations.

Primary path : pypandoc -> pandoc --mathml
               Converts the generated Markdown into a .docx. The --mathml
               flag makes pandoc emit native Word (OMML) equations so the
               LaTeX is rendered by Word itself, not shown as raw text.

Fallback    : python-docx + latex2mathml
               Used when pandoc/pypandoc is unavailable (e.g. Streamlit
               Cloud). LaTeX is converted to MathML with latex2mathml and
               then into Word's native OMML, which is injected into the
               document XML.
"""

import os
import re

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_XML_NS = "http://www.w3.org/XML/1998/namespace"

# Matches $$display$$, **bold**, *italic* and $latex$ inline spans.
_INLINE_RE = re.compile(
    r'(\$\$[^$]+\$\$|\*\*[^*\n]+\*\*|\*[^*\n]+\*|\$[^$\n]+\$)'
)

# Optional per-variation solution fields generated from user instructions.
_SOLUTION_TITLES = {
    "solution_by_concept": "Penyelesaian (Konsep Dasar)",
    "solution_by_trick": "Penyelesaian (Cara Cepat/Trik)",
}


def _sanitize_option(text):
    """Collapse newlines inside a single option so it stays on one line."""
    return re.sub(r"\s+", " ", text or "").strip()


def build_markdown(questions):
    """Build the full Markdown document for the given list of questions."""
    lines = ["# Bank Soal & Variasi Matematika", ""]

    for idx, item in enumerate(questions, 1):
        original = item.get("original", {})
        variations = item.get("variations", {})
        page = item.get("page")

        heading = f"## Soal {idx}"
        if page:
            heading += f" (Halaman {page})"
        lines += [heading, "", f"**ID:** {original.get('id', 'Unknown')}", ""]

        lines += ["### Soal Asli", "", original.get("question_text", "").strip(), ""]

        options = original.get("options") or []
        if options:
            lines += ["**Opsi Jawaban:**", ""]
            for i, opt in enumerate(options):
                lines.append(f"- **{chr(65 + i)}.** {_sanitize_option(opt)}")
            lines.append("")

        for variant in ("easier", "harder"):
            label = "Mudah" if variant == "easier" else "Sulit"
            if variant not in variations:
                continue
            vdata = variations[variant]
            lines += [f"### Variasi Lebih {label}", "",
                      vdata.get("question_text", "").strip(), ""]
            vopts = vdata.get("options") or []
            for i, opt in enumerate(vopts):
                lines.append(f"- **{chr(65 + i)}.** {_sanitize_option(opt)}")
            lines.append("")
            for key, title in _SOLUTION_TITLES.items():
                solution = vdata.get(key)
                if solution and solution.strip():
                    lines += [f"**{title}:**", "", solution.strip(), ""]

        lines += ["---", ""]

    return "\n".join(lines).strip()


def export_docx(questions, output_path):
    """Export the questions list to a professional Word document."""
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

    try:
        import pypandoc

        markdown_text = build_markdown(questions)
        print("  [3/4] Building DOCX via pandoc --mathml (native Word equations)...")
        pypandoc.convert_text(
            markdown_text,
            "docx",
            format="markdown",
            extra_args=["--mathml"],
            outputfile=output_path,
        )
    except Exception as exc:
        print(f"  ⚠ pypandoc/pandoc unavailable ({exc}); "
              f"falling back to python-docx + latex2mathml...")
        _export_with_python_docx(questions, output_path)

    return output_path


# --------------------------------------------------------------------------- #
# Fallback: python-docx + latex2mathml (MathML -> OMML injected into the XML) #
# --------------------------------------------------------------------------- #

def _split_inline(text):
    """Split text into math, bold and italic segments."""
    parts = []
    for chunk in _INLINE_RE.split(text):
        if not chunk:
            continue
        if chunk.startswith("$$") and chunk.endswith("$$") and len(chunk) > 4:
            parts.append({"type": "math", "text": chunk[2:-2]})
        elif chunk.startswith("$") and chunk.endswith("$") and len(chunk) > 2:
            parts.append({"type": "math", "text": chunk[1:-1]})
        elif chunk.startswith("**") and chunk.endswith("**") and len(chunk) > 4:
            parts.append({"type": "run", "text": chunk[2:-2],
                          "bold": True, "italic": False})
        elif chunk.startswith("*") and chunk.endswith("*") and len(chunk) > 2:
            parts.append({"type": "run", "text": chunk[1:-1],
                          "bold": False, "italic": True})
        else:
            parts.append({"type": "run", "text": chunk,
                          "bold": False, "italic": False})
    return parts


def _omml_from_mathml(mathml_xml):
    """Convert a latex2mathml MathML string into a Word OMML <m:oMath> element."""
    from lxml import etree

    root = etree.fromstring(mathml_xml)

    def qm(tag):
        return f"{{{_M_NS}}}{tag}"

    def token(node):
        return etree.QName(node).localname in ("mi", "mn", "mo", "mtext")

    def children_of(node):
        """Leaf token nodes convert themselves; everything else is a container."""
        return [node] if token(node) else list(node)

    _DELIM_OPEN = {"[": "[", "(": "(", "{": "{", "|": "|"}
    _DELIM_CLOSE = {"]": "]", ")": ")", "}": "}", "|": "|"}

    def _is_delimiter(node):
        if etree.QName(node).localname != "mo":
            return False
        text = (node.text or "").strip()
        return text in _DELIM_OPEN or text in _DELIM_CLOSE

    def _delims_around(kids, idx):
        beg = end = ""
        if idx > 0 and _is_delimiter(kids[idx - 1]):
            beg = _DELIM_OPEN.get(kids[idx - 1].text.strip(), "")
        if idx + 1 < len(kids) and _is_delimiter(kids[idx + 1]):
            end = _DELIM_CLOSE.get(kids[idx + 1].text.strip(), "")
        return beg, end

    def _convert_matrix(parent, node, delims=("", "")):
        beg, end = delims
        matrix = etree.SubElement(parent, qm("m"))
        mpr = etree.SubElement(matrix, qm("mPr"))
        if beg:
            el = etree.SubElement(mpr, qm("begChr"))
            el.set(qm("val"), beg)
        if end:
            el = etree.SubElement(mpr, qm("endChr"))
            el.set(qm("val"), end)
        mcs = etree.SubElement(mpr, qm("mcs"))
        mc = etree.SubElement(mcs, qm("mc"))
        mcpr = etree.SubElement(mc, qm("mcPr"))
        first_row = next(
            (r for r in node if etree.QName(r).localname == "mtr"), None)
        ncols = (sum(1 for c in first_row
                     if etree.QName(c).localname == "mtd")
                 if first_row is not None else 1)
        count = etree.SubElement(mcpr, qm("count"))
        count.set(qm("val"), str(ncols))
        jc = etree.SubElement(mcpr, qm("mcJc"))
        jc.set(qm("val"), "center")
        for mtr in node:
            if etree.QName(mtr).localname != "mtr":
                continue
            row = etree.SubElement(matrix, qm("mr"))
            for mtd in mtr:
                if etree.QName(mtd).localname != "mtd":
                    continue
                cell = etree.SubElement(row, qm("e"))
                for child in children_of(mtd):
                    convert(cell, child)

    def convert(parent, node):
        tag = etree.QName(node).localname
        if token(node):
            text = node.text or ""
            if text:
                run = etree.SubElement(parent, qm("r"))
                t = etree.SubElement(run, qm("t"))
                t.set(f"{{{_XML_NS}}}space", "preserve")
                t.text = text
        elif tag == "mrow":
            kids = list(node)
            table_idx = [i for i, c in enumerate(kids)
                         if etree.QName(c).localname == "mtable"]
            adjacent = set()
            for ti in table_idx:
                if ti > 0:
                    adjacent.add(ti - 1)
                if ti + 1 < len(kids):
                    adjacent.add(ti + 1)
            for i, child in enumerate(kids):
                ctag = etree.QName(child).localname
                if i in adjacent and _is_delimiter(child):
                    continue
                if ctag == "mtable":
                    _convert_matrix(parent, child, _delims_around(kids, i))
                else:
                    convert(parent, child)
        elif tag == "mtable":
            _convert_matrix(parent, node)
        elif tag == "mfrac":
            frac = etree.SubElement(parent, qm("f"))
            num = etree.SubElement(frac, qm("num"))
            den = etree.SubElement(frac, qm("den"))
            if len(node) > 0:
                for child in children_of(node[0]):
                    convert(num, child)
            if len(node) > 1:
                for child in children_of(node[1]):
                    convert(den, child)
        elif tag == "msqrt":
            srad = etree.SubElement(parent, qm("sRad"))
            etree.SubElement(srad, qm("deg"))
            e = etree.SubElement(srad, qm("e"))
            for child in children_of(node):
                convert(e, child)
        elif tag in ("msup", "msub", "msubsup"):
            if tag == "msup":
                wrapper = etree.SubElement(parent, qm("sSup"))
                sub_tag, sup_tag = None, "sup"
            elif tag == "msub":
                wrapper = etree.SubElement(parent, qm("sSub"))
                sub_tag, sup_tag = "sub", None
            else:
                wrapper = etree.SubElement(parent, qm("sSubSup"))
                sub_tag, sup_tag = "sub", "sup"
            base = etree.SubElement(wrapper, qm("e"))
            for child in children_of(node[0]):
                convert(base, child)
            if sub_tag:
                sub = etree.SubElement(wrapper, qm(sub_tag))
                for child in children_of(node[1]):
                    convert(sub, child)
            if sup_tag:
                sup = etree.SubElement(wrapper, qm(sup_tag))
                idx = 1 if not sub_tag else 2
                for child in children_of(node[idx]):
                    convert(sup, child)
        else:
            for child in children_of(node):
                convert(parent, child)

    omath = etree.Element(qm("oMath"))
    for child in root:
        convert(omath, child)
    return omath


def _append_math(paragraph, latex, latex_to_mathml):
    """Append a LaTeX formula as a native Word equation (OMML)."""
    if latex_to_mathml is not None:
        try:
            mathml = latex_to_mathml(latex)
            paragraph._p.append(_omml_from_mathml(mathml))
            return
        except Exception:
            pass
    run = paragraph.add_run(f"${latex}$")
    run.italic = True


def _add_rich_paragraph(doc, text, latex_to_mathml, bullet=False):
    """Add a paragraph, honouring inline markdown and $...$ LaTeX."""
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    for part in _split_inline(text):
        if part["type"] == "math":
            _append_math(p, part["text"], latex_to_mathml)
        else:
            run = p.add_run(part["text"])
            run.bold = part.get("bold", False)
            run.italic = part.get("italic", False)
    return p


def _export_with_python_docx(questions, output_path):
    """Fallback exporter built directly on python-docx + latex2mathml."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    try:
        from latex2mathml.converter import convert as latex_to_mathml
    except ImportError:
        latex_to_mathml = None

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_heading("Bank Soal & Variasi Matematika", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, item in enumerate(questions, 1):
        original = item.get("original", {})
        variations = item.get("variations", {})
        page = item.get("page")

        heading = f"Soal {idx}"
        if page:
            heading += f" (Halaman {page})"
        doc.add_heading(heading, level=2)

        _add_rich_paragraph(doc, f"**ID:** {original.get('id', 'Unknown')}",
                            latex_to_mathml)

        doc.add_heading("Soal Asli", level=3)
        _add_rich_paragraph(doc, original.get("question_text", ""),
                            latex_to_mathml)

        options = original.get("options") or []
        if options:
            doc.add_heading("Opsi Jawaban", level=4)
            for i, opt in enumerate(options):
                _add_rich_paragraph(
                    doc, f"**{chr(65 + i)}.** {_sanitize_option(opt)}",
                    latex_to_mathml, bullet=True)

        for variant in ("easier", "harder"):
            label = "Mudah" if variant == "easier" else "Sulit"
            if variant not in variations:
                continue
            doc.add_heading(f"Variasi Lebih {label}", level=3)
            _add_rich_paragraph(
                doc, variations[variant].get("question_text", ""),
                latex_to_mathml)
            vopts = variations[variant].get("options") or []
            for i, opt in enumerate(vopts):
                _add_rich_paragraph(
                    doc, f"**{chr(65 + i)}.** {_sanitize_option(opt)}",
                    latex_to_mathml, bullet=True)
            for key, title in _SOLUTION_TITLES.items():
                solution = variations[variant].get(key)
                if solution and solution.strip():
                    _add_rich_paragraph(
                        doc, f"**{title}:** {solution.strip()}",
                        latex_to_mathml)

    doc.save(output_path)
    print(f"  ✅ DOCX built with python-docx + latex2mathml: {output_path}")

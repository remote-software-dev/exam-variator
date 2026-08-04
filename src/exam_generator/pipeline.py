import os
import sys
import json
import base64
import re
import litellm
from dotenv import load_dotenv

# Add the project root to the path so we can import modules
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../..")))

load_dotenv()

# LiteLLM model fallback chain: tries models in order until one succeeds
EXTRACTION_MODELS = [
    "groq/llama-3.2-11b-vision-preview",
    "groq/llama-3.3-70b-versatile",
    "openai/gpt-4o-mini",
]

VARIATION_MODELS = [
    "groq/llama-3.3-70b-versatile",
    "openai/gpt-4o-mini",
]

def encode_image(image_path):
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode('utf-8')

def _extract_json(text):
    """Robustly extract the first JSON object from model output."""
    match = re.search(r'```json\s*(\{.*?\})\s*```', text, re.DOTALL)
    if not match:
        match = re.search(r'(\{.*\})', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1).replace("'", '"'))
        except json.JSONDecodeError:
            return None
    return None

def extract_question_from_image(image_path):
    print("  [1/4] Extracting question from image via LiteLLM (with fallbacks)...")
    base64_image = encode_image(image_path)

    system_prompt = (
        "You are an expert at extracting Indonesian math exam questions from scanned images.\n"
        "Return ONLY a valid JSON object with keys: 'id', 'question_text', 'options' (list of strings).\n\n"
        "RULES:\n"
        "- Use LaTeX math notation enclosed in $ delimiters for all formulas "
        "(e.g., $\\frac{a}{b}$, $x^2$, $\\sqrt{3}$).\n"
        "- 'id' must be the alphanumeric ID printed on the paper (e.g., '25MATBLGBRLM01SU-000000-0246').\n"
        "  If no ID is visible or the ID is just a number like '1', generate a unique one: "
        "'EXAM-<RANDOM8HEX>'.\n"
        "- Handle ALL question formats:\n"
        "  * Standard multiple choice (A/B/C/D/E) → put options in 'options' list.\n"
        "  * Benar/Salah (True/False) tables → 'options' should be a list of statements "
        "each prefixed with the table label, e.g. ['Pernyataan 1: Benar', 'Pernyataan 2: Salah'].\n"
        "  * Multi-part / stem questions (e.g., 'pernyataan (1), (2), (3)') → put each "
        "statement as a separate item in 'options'.\n"
        "- Preserve the original Indonesian language.\n"
        "- Use double quotes for all JSON keys and string values."
    )

    user_content = [
        {"type": "text", "text": "Extract the first complete question from this image as JSON."},
        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{base64_image}"}},
    ]

    for model in EXTRACTION_MODELS:
        try:
            response = litellm.completion(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_content},
                ],
                temperature=0.1,
            )
            raw = response.choices[0].message.content
            result = _extract_json(raw)
            if result:
                return result
            print(f"  ⚠ {model} returned unparseable JSON, trying next...")
        except Exception as e:
            print(f"  ⚠ {model} failed: {e}, trying next...")

    print("  ❌ All extraction models failed.")
    return None

def generate_variations(original_q):
    print("  [2/4] Generating easier and harder variations via LiteLLM (with fallbacks)...")

    system_prompt = (
        "You are a math exam question writer for Indonesian secondary schools.\n"
        "Given an original multiple-choice question, produce two variations: 'easier' and 'harder'.\n\n"
        "STRICT FORMAT RULES:\n"
        "- Return ONLY a valid JSON object.\n"
        "- Top-level keys: 'easier' and 'harder'.\n"
        "- Each variation must contain:\n"
        "  * 'question_text': string (the question stem)\n"
        "  * 'options': array of exactly 5 strings (labeled A through E in the output)\n"
        "- Do NOT convert multiple-choice into essay or fill-in-the-blank questions.\n"
        "- Do NOT change the number of options. Every variation MUST have exactly 5 options.\n"
        "- Do NOT include option labels (A., B., etc.) inside the option strings — just the answer text.\n"
        "- Use LaTeX math notation enclosed in $ delimiters for all formulas.\n"
        "- Keep the same mathematical topic and difficulty relative to the label (easier = simpler numbers/steps, "
        "harder = more complex numbers/steps or additional concepts).\n"
        "- Preserve the original Indonesian language.\n"
        "- Use double quotes for all JSON keys and string values.\n"
        "- Do NOT wrap the JSON in markdown code fences."
    )

    user_prompt = (
        f"Original question:\n{json.dumps(original_q, ensure_ascii=False, indent=2)}\n\n"
        "Generate the 'easier' and 'harder' variations now."
    )

    for model in VARIATION_MODELS:
        try:
            response = litellm.completion(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=0.7,
            )
            raw = response.choices[0].message.content
            result = _extract_json(raw)
            if result and "easier" in result and "harder" in result:
                return result
            print(f"  ⚠ {model} returned invalid variation structure, trying next...")
        except Exception as e:
            print(f"  ⚠ {model} failed: {e}, trying next...")

    print("  ❌ All variation models failed.")
    return None

def generate_docx(data, output_path):
    print("  [3/4] Building DOCX with python-docx...")
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title = doc.add_heading("Bank Soal & Variasi Matematika", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading(f"Soal Asli (ID: {data['original']['id']})", level=2)
    doc.add_paragraph(data['original']['question_text'])
    for i, opt in enumerate(data['original']['options']):
        doc.add_paragraph(f"{chr(65+i)}. {opt}", style='List Number')

    for variant in ['easier', 'harder']:
        if variant in data['variations']:
            label = 'Mudah' if variant == 'easier' else 'Sulit'
            doc.add_heading(f"Variasi Lebih {label}", level=2)
            doc.add_paragraph(data['variations'][variant]['question_text'])
            for i, opt in enumerate(data['variations'][variant]['options']):
                doc.add_paragraph(f"{chr(65+i)}. {opt}", style='List Number')

    doc.save(output_path)
    print(f"  [4/4] DOCX saved to: {output_path}")
    print(f"✅ Success! Output saved to: {output_path}")

def run_pipeline(pdf_path, output_docx):
    """Run the full pipeline: PDF -> PNG -> Extract -> Vary -> DOCX."""
    import fitz  # PyMuPDF

    print("🚀 Starting End-to-End Exam Generator Pipeline...\n")

    # Render first page of PDF to PNG
    pages_dir = "data/outputs/pages"
    os.makedirs(pages_dir, exist_ok=True)
    doc = fitz.open(pdf_path)
    page = doc[0]
    pix = page.get_pixmap(dpi=200)
    page_png = os.path.join(pages_dir, "page_01.png")
    pix.save(page_png)
    doc.close()
    print(f"  ✅ Rendered page 1 to {page_png}")

    # 1. Extract
    original_q = extract_question_from_image(page_png)
    if not original_q:
        raise RuntimeError("Failed to extract question from image.")
    print(f"  ✅ Extracted: {original_q.get('id', 'Unknown ID')}")

    # 2. Vary
    variations = generate_variations(original_q)
    if not variations:
        raise RuntimeError("Failed to generate variations.")
    print("  ✅ Variations generated.")

    # 3. Export
    os.makedirs(os.path.dirname(output_docx), exist_ok=True)
    generate_docx({
        "original": original_q,
        "variations": variations
    }, output_docx)

    return output_docx


def main():
    run_pipeline(
        pdf_path="data/inputs/SOAL TKA Matematika SMA 2025 Umum.pdf",
        output_docx="data/outputs/final_pipeline_test.docx",
    )

if __name__ == "__main__":
    main()
